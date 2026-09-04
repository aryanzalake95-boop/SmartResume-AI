import os
import re


def _clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _pdf_pypdf(path):
    from PyPDF2 import PdfReader
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n".join(parts)


def _pdf_pymupdf(path):
    try:
        import fitz
    except ImportError:
        return ""
    doc = fitz.open(path)
    try:
        return "\n".join(page.get_text("text") or "" for page in doc)
    finally:
        doc.close()


def _docx_text(path):
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_resume_text(path, ext):
    ext = ext.lower().lstrip(".")
    if not os.path.exists(path):
        raise FileNotFoundError("Uploaded resume could not be found.")
    if ext == "pdf":
        text = _pdf_pypdf(path)
        if len(re.sub(r"\s+", "", text)) < 80:
            fallback = _pdf_pymupdf(path)
            if len(re.sub(r"\s+", "", fallback)) > len(re.sub(r"\s+", "", text)):
                text = fallback
        text = _clean_text(text)
        if len(re.sub(r"\s+", "", text)) < 30:
            raise ValueError("This PDF contains no readable text. It may be a scanned/image-only PDF. Please upload a text-based PDF or DOCX resume.")
        return text
    if ext == "docx":
        text = _clean_text(_docx_text(path))
        if len(re.sub(r"\s+", "", text)) < 30:
            raise ValueError("The DOCX file does not contain enough readable resume text.")
        return text
    if ext == "txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = _clean_text(f.read())
        if len(re.sub(r"\s+", "", text)) < 30:
            raise ValueError("The TXT file does not contain enough readable resume text.")
        return text
    raise ValueError("Unsupported file type.")
